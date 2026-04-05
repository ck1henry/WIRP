"""
Generates WIRP_Methodology.docx — one-page methodology note.
Run once: python make_methodology_doc.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_PATH   = os.path.join(SCRIPT_DIR, "WIRP_Methodology.docx")

# ── Colour palette ──────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0f, 0x29, 0x42)
AMBER  = RGBColor(0xb4, 0x53, 0x09)
SLATE  = RGBColor(0x47, 0x56, 0x69)
BLACK  = RGBColor(0x1e, 0x29, 0x3b)
GREEN  = RGBColor(0x15, 0x80, 0x3d)
LGREY  = RGBColor(0xf1, 0xf5, 0xf9)

def set_cell_bg(cell, hex_str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_str)
    tcPr.append(shd)

def add_run(para, text, bold=False, italic=False, size=10, color=None):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

doc = Document()

# ── Page margins (narrow) ───────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

# ── Default style ───────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(9.5)
style.font.color.rgb = BLACK

# ════════════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ════════════════════════════════════════════════════════════════════════════
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_run(title_para, 'WIRP  —  World Interest Rate Probabilities',
        bold=True, size=16, color=NAVY)

sub_para = doc.add_paragraph()
sub_para.paragraph_format.space_before = Pt(2)
sub_para.paragraph_format.space_after  = Pt(10)
add_run(sub_para,
        'Methodology Note  ·  G10 Core  ·  5 Central Banks  ·  Real OIS / Futures Data',
        italic=True, size=9, color=SLATE)

# ── Thin amber rule under title ─────────────────────────────────────────────
rule = doc.add_paragraph()
rule.paragraph_format.space_before = Pt(0)
rule.paragraph_format.space_after  = Pt(10)
pPr  = rule._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'),   'single')
bottom.set(qn('w:sz'),    '6')
bottom.set(qn('w:space'), '1')
bottom.set(qn('w:color'), 'B45309')
pBdr.append(bottom)
pPr.append(pBdr)

# ════════════════════════════════════════════════════════════════════════════
# SECTION HELPER
# ════════════════════════════════════════════════════════════════════════════
def section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    add_run(p, text.upper(), bold=True, size=8.5, color=AMBER)
    return p

def body_para(doc, text, space_after=4):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    for run in p.runs:
        run.font.size = Pt(9.5)
        run.font.color.rgb = BLACK
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.4)
    add_run(p, text, size=9.5, color=BLACK)
    return p

# ════════════════════════════════════════════════════════════════════════════
# 1. PURPOSE
# ════════════════════════════════════════════════════════════════════════════
section_heading(doc, '1.  Purpose')
body_para(doc,
    'WIRP (World Interest Rate Probabilities) is a real-time dashboard that replicates the '
    'analytical logic of the Bloomberg Terminal WIRP function. It extracts market-implied '
    'probability distributions over future central bank policy rate decisions from overnight '
    'indexed swap (OIS) curves and short-term interest rate futures, and presents them in an '
    'interactive, single-file HTML application.')

# ════════════════════════════════════════════════════════════════════════════
# 2. COVERAGE
# ════════════════════════════════════════════════════════════════════════════
section_heading(doc, '2.  Coverage')
body_para(doc,
    'Five central banks are covered, each with real market-derived implied rates. Meeting dates '
    'and current policy rates are sourced as follows:', space_after=3)

tbl = doc.add_table(rows=1, cols=4)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
tbl.style = 'Table Grid'

hdr_cells = tbl.rows[0].cells
for cell, txt in zip(hdr_cells, ['Region', 'Bank', 'Abbr.', 'Implied Rate Source']):
    set_cell_bg(cell, '0F2942')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, txt, bold=True, size=8.5, color=RGBColor(0xff,0xff,0xff))

rows = [
    ('Americas', 'Federal Reserve',          'FOMC',   'CME 30-day Fed Funds futures via TradingView scanner (Yahoo Finance fallback); base rate = EFFR (FRED DFF)'),
    ('Americas', 'Bank of Canada',           'BOC',    'BOC Valet API — CORRA overnight + 2Y gov\'t bond'),
    ('Europe',   'European Central Bank',    'ECB GC', 'Eurex 3M EURIBOR futures (FEU3) via TradingView, DFR-anchored; ECB SDW YC fallback'),
    ('Europe',   'Bank of England',          'MPC',    'BOE SONIA OIS instantaneous forward curve'),
    ('Asia-Pac', 'Reserve Bank of Australia','RBA',    'RBA F01 BABs/NCDs (1M/3M/6M), BBSW-OIS adjusted'),
]
for i, (region, name, abbr, src) in enumerate(rows):
    r = tbl.add_row()
    bg = 'F8FAFC' if i % 2 == 0 else 'FFFFFF'
    for cell, txt in zip(r.cells, [region, name, abbr, src]):
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        add_run(p, txt, size=8.5, color=BLACK)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ════════════════════════════════════════════════════════════════════════════
# 3. METHODOLOGY
# ════════════════════════════════════════════════════════════════════════════
section_heading(doc, '3.  Implied Rate Methodology')
body_para(doc,
    'For each upcoming central bank meeting, the tool derives an OIS-implied policy rate and '
    'computes hike/cut probabilities using the following approach:')

bullet(doc,
    'US (FOMC):  CME 30-day Fed Funds futures (ZQ) fetched via TradingView scanner API '
    '(CBOT:ZQ{M}{YYYY}, field "close").  For a meeting in month M the month M+1 contract '
    'captures the post-decision rate — mirrors CME FedWatch methodology.  '
    'implied_rate = 100 − price.  Yahoo Finance serves as per-contract fallback.')
bullet(doc,
    'EU (ECB):  Eurex 3-month EURIBOR futures (FEU3{M}{YYYY}) fetched via TradingView.  '
    'implied_euribor = 100 − price.  A EURIBOR-OIS spread is calibrated dynamically: '
    'spread = nearest_contract_implied_euribor − current_DFR.  '
    'ECB implied = implied_euribor − spread, so the curve is anchored to DFR at the near end '
    'and preserves the market-implied shape for subsequent meetings.  '
    'Falls back to the ECB SDW AAA government bond yield curve if futures are unavailable.')
bullet(doc,
    'UK (MPC):  BOE SONIA OIS instantaneous forward curve (ZIP file, 60 monthly tenors).  '
    'Interpolated to each meeting date.')
bullet(doc,
    'CA (BOC):  Two-point linear interpolation between BOC Valet API overnight CORRA and '
    '2-year government benchmark bond yield.')
bullet(doc,
    'AU (RBA):  RBA F01 BABs/NCDs at 1M/3M/6M maturities, adjusted by a fixed BBSW-OIS spread.')
bullet(doc,
    'Implied Change (bps):  Δ = Implied Rate − Base Rate, expressed in basis points.  '
    'For the US, the base rate is the Effective Fed Funds Rate (EFFR, FRED series DFF) — '
    'the same basis Bloomberg WIRP uses — not the target-range midpoint.')
bullet(doc,
    'Step size of 25 bps is assumed throughout (standard for G10 central banks).')

body_para(doc,
    'Colour convention:  positive Δ (hikes priced) shown in green; negative Δ (cuts priced) '
    'shown in red; exactly zero in grey.',
    space_after=6)

# ════════════════════════════════════════════════════════════════════════════
# 4. MARKET STANCE CLASSIFICATION
# ════════════════════════════════════════════════════════════════════════════
section_heading(doc, '4.  Market Stance Classification')
body_para(doc,
    'Each central bank is assigned a market stance badge based on the cumulative implied '
    'change across all displayed meetings:', space_after=3)

stance_tbl = doc.add_table(rows=1, cols=3)
stance_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
stance_tbl.style = 'Table Grid'
for cell, txt in zip(stance_tbl.rows[0].cells, ['Stance', 'Threshold (cumulative bps)', 'Interpretation']):
    set_cell_bg(cell, '0F2942')
    add_run(cell.paragraphs[0], txt, bold=True, size=8.5, color=RGBColor(0xff,0xff,0xff))

for stance, thresh, interp in [
    ('Very Dovish', '≤ −50 bps',           'Markets price ≥2 cuts within the meeting window'),
    ('Dovish',      '−49 to −15 bps',      'Markets lean toward one or more cuts'),
    ('Neutral',     '−14 to +14 bps',      'No meaningful directional tilt priced in'),
    ('Hawkish',     '≥ +15 bps',           'Markets lean toward one or more hikes'),
]:
    r = stance_tbl.add_row()
    for cell, txt in zip(r.cells, [stance, thresh, interp]):
        add_run(cell.paragraphs[0], txt, size=8.5, color=BLACK)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ════════════════════════════════════════════════════════════════════════════
# 5. MEETING REPRICING HISTORY
# ════════════════════════════════════════════════════════════════════════════
section_heading(doc, '5.  Meeting Repricing History')
body_para(doc,
    'The repricing history chart shows how OIS-implied expectations for a specific future '
    'meeting have evolved over the preceding 12 months.  The methodology is as follows:')
bullet(doc,
    'For each month t going back up to 12 months, the estimated implied rate for meeting M '
    'is: Implied(t) = ImpliedToday + α × (SpotThen − SpotNow) + ε(t)')
bullet(doc,
    'SpotThen is estimated by reverse-extrapolating the current policy rate via the drift '
    'coefficient: SpotThen ≈ Spot − Drift × (t / 12) × 1.35')
bullet(doc,
    'α = 0.72 captures how tightly forward meeting pricing tracks the evolving spot rate. '
    'ε(t) is deterministic pseudo-noise, stable across page loads, that increases with t to '
    'reflect greater uncertainty further back in time.')
body_para(doc,
    'In production, this panel would be populated from a time-series database of daily OIS '
    'snapshots, showing the actual evolution of market pricing for each specific meeting date.',
    space_after=6)

# ════════════════════════════════════════════════════════════════════════════
# 6. DATA PIPELINE
# ════════════════════════════════════════════════════════════════════════════
section_heading(doc, '6.  Data Pipeline')
bullet(doc,
    'update_data.py fetches live policy rates, meeting calendars, and implied rates for all '
    'five markets and injects results into wirp.html between '
    '<!-- WIRP_DATA_BEGIN --> and <!-- WIRP_DATA_END --> markers.')
bullet(doc,
    'Primary implied-rate sources use the TradingView scanner API (no key required): '
    'CBOT:ZQ{M}{YYYY} for US Fed Funds futures and EUREX:FEU3{M}{YYYY} for ECB EURIBOR futures.  '
    'Yahoo Finance is the per-contract fallback for US; ECB SDW YC is the fallback for EU.')
bullet(doc,
    'US base rate is fetched from FRED API (series DFF — daily EFFR), matching Bloomberg\'s '
    'WIRP methodology.  Falls back to the Fed target-range midpoint if FRED is unavailable.')
bullet(doc,
    'Data is also persisted to wirp_data.json; per-market graceful degradation ensures a fetch '
    'failure for one bank does not block the others.')
bullet(doc,
    'Automation is via cron on an Oracle Cloud VM (hourly, 0 * * * *).  '
    'The script logs to wirp_update.log.')

# ════════════════════════════════════════════════════════════════════════════
# 7. LIMITATIONS
# ════════════════════════════════════════════════════════════════════════════
section_heading(doc, '7.  Limitations & Assumptions')
bullet(doc, 'No convexity adjustment is applied to OIS or futures-implied rates.')
bullet(doc, 'Probability calculations assume a binary 25 bps step; multi-step pricing (e.g. 50 bps) is not modelled.')
bullet(doc, 'ECB implied rates use Eurex EURIBOR futures anchored to the DFR.  The EURIBOR-OIS '
    'spread is calibrated daily to the nearest available contract; its level and term structure '
    'vary over time and are not explicitly modelled.')
bullet(doc, 'For ECB meetings within the same calendar month as the nearest EURIBOR contract, '
    'the implied rate equals the DFR by construction (the anchor point); no independent market '
    'signal is available for that meeting from this source.')
bullet(doc, 'US ±bps figures are relative to the EFFR (daily fixing), which trades within the '
    'Fed\'s target range and may differ slightly from the target midpoint.')
bullet(doc, 'TradingView scanner API is an unofficial endpoint; availability and field names '
    'may change without notice.  Yahoo Finance (US) and ECB SDW YC (EU) are maintained as fallbacks.')
bullet(doc, 'CA implied rates are interpolated from two points only (overnight CORRA + 2Y bond); the curve shape is linear.')
bullet(doc, 'AU implied rates use BABs/NCDs with a fixed BBSW-OIS spread adjustment; the spread varies over time.')
bullet(doc, 'Historical repricing series are simulated; they are illustrative and not based on live OIS time-series data.')

# ════════════════════════════════════════════════════════════════════════════
# FOOTER LINE
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr2  = footer_para._p.get_or_add_pPr()
pBdr2 = OxmlElement('w:pBdr')
top   = OxmlElement('w:top')
top.set(qn('w:val'),   'single')
top.set(qn('w:sz'),    '4')
top.set(qn('w:space'), '1')
top.set(qn('w:color'), 'CBD5E1')
pBdr2.append(top)
pPr2.append(pBdr2)
add_run(footer_para,
        'WIRP v2.1  ·  For internal research use only  ·  OIS / Futures-implied estimates; not investment advice',
        italic=True, size=7.5, color=SLATE)

doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
